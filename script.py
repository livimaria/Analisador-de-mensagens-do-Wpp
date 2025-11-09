import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import re
import io
from datetime import datetime
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4

# -------- Função para classificar mensagens --------
def classificar_mensagem(texto):
    texto = texto.lower()
    if any(palavra in texto for palavra in ["consulta", "médico", "retorno"]):
        return "Consulta"
    elif any(palavra in texto for palavra in ["exame", "resultado", "teste"]):
        return "Exame"
    elif any(palavra in texto for palavra in ["cirurgia", "procedimento", "operação"]):
        return "Cirurgia"
    else:
        return "Outros"

# -------- Função para processar arquivos WhatsApp --------
def processar_arquivos(arquivos):
    todas_mensagens = []

    for arquivo in arquivos:
        conteudo = arquivo.read().decode("utf-8", errors="ignore")
        linhas = conteudo.splitlines()

        for linha in linhas:
            linha = linha.strip()

            # Novo parser genérico (compatível com Android e iPhone)
            # Exemplo: [10/10/2025, 09:35] João: Olá
            if "] " in linha and ":" in linha:
                try:
                    parte_data, resto = linha.split("] ", 1)
                    parte_data = parte_data.replace("[", "").strip()

                    # Divide remetente e mensagem
                    if ": " in resto:
                        nome, mensagem = resto.split(": ", 1)
                        data_str = parte_data.split(",")[0].strip()
                        hora_str = parte_data.split(",")[1].strip() if "," in parte_data else ""

                        categoria = classificar_mensagem(mensagem)
                        todas_mensagens.append([data_str, hora_str, nome.strip(), mensagem.strip(), categoria])
                except Exception:
                    continue  # ignora linhas que não seguem o formato esperado

    df = pd.DataFrame(todas_mensagens, columns=["Data", "Hora", "Remetente", "Mensagem", "Categoria"])
    return df

# -------- Função para gerar gráficos --------
import plotly.express as px  # adicione no topo do arquivo

def gerar_graficos(df):
    st.subheader("📈 Quantidade de Mensagens por Categoria")

    # Contagem por categoria
    contagem = df['Categoria'].value_counts().reset_index()
    contagem.columns = ["Categoria", "Quantidade"]

    # Paleta de cores personalizada
    cores = {
        "Consulta": "#3B82F6",  # Azul
        "Exame": "#10B981",     # Verde
        "Cirurgia": "#EF4444",  # Vermelho
        "Outros": "#9CA3AF"     # Cinza
    }

    # Gráfico de barras interativo (Plotly)
    fig = px.bar(
        contagem,
        x="Categoria",
        y="Quantidade",
        color="Categoria",
        color_discrete_map=cores,
        text="Quantidade",
        title="Total de Mensagens por Categoria"
    )

    fig.update_traces(textposition="outside")
    fig.update_layout(
        xaxis_title="Categoria",
        yaxis_title="Quantidade de Mensagens",
        showlegend=False,
        plot_bgcolor="rgba(0,0,0,0)",
        yaxis=dict(showgrid=True)
    )

    st.plotly_chart(fig, use_container_width=True)

# -------- Funções de exportação --------
def exportar_excel(df):
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Relatório')
    return buffer

def exportar_word(df):
    doc = Document()
    doc.add_heading("Relatório de Conversas WhatsApp", level=1)
    doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    tabela = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = tabela.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col

    for _, row in df.iterrows():
        linha = tabela.add_row().cells
        for i, val in enumerate(row):
            linha[i].text = str(val)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def exportar_pdf(df):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    elementos = [Paragraph("Relatório de Conversas WhatsApp", styles['Title'])]
    elementos.append(Spacer(1, 12))
    elementos.append(Paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles['Normal']))
    elementos.append(Spacer(1, 12))

    resumo = df['Categoria'].value_counts()
    for cat, qtd in resumo.items():
        elementos.append(Paragraph(f"{cat}: {qtd} mensagens", styles['Normal']))

    doc.build(elementos)
    buffer.seek(0)
    return buffer

# -------- Interface principal --------
st.title("Analisador Inteligente de Conversas do WhatsApp - Ecardio")
st.markdown("Carregue **múltiplos arquivos .txt** exportados do WhatsApp para gerar o relatório semanal.")

arquivos = st.file_uploader("Selecione os arquivos", type=["txt"], accept_multiple_files=True)

if arquivos:
    with st.spinner("Processando arquivos..."):
        df = processar_arquivos(arquivos)
        st.success(f"{len(df)} mensagens processadas com sucesso!")

        gerar_graficos(df)

        st.subheader("📋 Dados Processados")
        st.dataframe(df)

        st.subheader("📤 Exportar Relatório")
        col1, col2, col3 = st.columns(3)

        with col1:
            st.download_button(
                "⬇️ Baixar Excel",
                data=exportar_excel(df),
                file_name="relatorio_whatsapp.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        with col2:
            st.download_button(
                "⬇️ Baixar Word",
                data=exportar_word(df),
                file_name="relatorio_whatsapp.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        with col3:
            st.download_button(
                "⬇️ Baixar PDF",
                data=exportar_pdf(df),
                file_name="relatorio_whatsapp.pdf",
                mime="application/pdf"
            )

else:
    st.info("Selecione um ou mais arquivos de conversas para iniciar a análise.")
