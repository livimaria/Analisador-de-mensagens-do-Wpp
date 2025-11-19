# src/export.py
from docx import Document
from docx.shared import Pt, Inches
import matplotlib
matplotlib.use("Agg")  # backend não interativo, bom para servidores
import matplotlib.pyplot as plt

def gerar_grafico_matplotlib(df):
    """
    Gera gráfico de barras usando matplotlib e salva em 'grafico_temp.png'.
    Retorna o caminho do arquivo.
    """
    contagem = df["categoria"].value_counts()
    categorias = contagem.index.tolist()
    quantidades = contagem.values.tolist()

    # cores correspondentes (mesmas do design)
    color_map = {
        "Consulta": "#1f77b4",
        "Exame": "#2ca02c",
        "Cirurgia": "#d62728",
        "Outros": "#7f7f7f"
    }
    cores = [color_map.get(cat, "#7f7f7f") for cat in categorias]

    fig, ax = plt.subplots(figsize=(8, 4.5))
    bars = ax.bar(categorias, quantidades, color=cores)
    ax.set_title("Distribuição de Mensagens por Categoria")
    ax.set_ylabel("Quantidade")
    ax.set_xlabel("Categoria")
    ax.bar_label(bars, padding=3)  # mostra os valores acima das barras
    plt.tight_layout()

    caminho = "grafico_temp.png"
    fig.savefig(caminho, dpi=150)
    plt.close(fig)
    return caminho


def export_to_word(df, filename="relatorio.docx", limite=200):
    """
    Gera um DOCX com:
     - resumo (total de mensagens / quantas incluídas)
     - gráfico (PNG gerado por matplotlib)
     - amostra das mensagens (limite)
    """
    doc = Document()

    # Título
    titulo = doc.add_heading("Relatório de Análise de Mensagens – Ecardio", level=1)
    try:
        titulo.runs[0].font.size = Pt(20)
    except Exception:
        pass

    # Resumo
    total_mensagens = len(df)
    mensagens_relatorio = min(limite, total_mensagens)

    texto_info = (
        f"Total de mensagens carregadas: {total_mensagens}\n"
        f"Mensagens incluídas neste relatório: {mensagens_relatorio}\n\n"
        "Atenção: o relatório contém apenas uma amostra, "
        "mas os gráficos e análises foram gerados com 100% dos dados."
    )
    doc.add_paragraph(texto_info)

    # Gráfico
    doc.add_heading("Gráfico de Distribuição por Categoria", level=2)
    caminho_img = gerar_grafico_matplotlib(df)
    doc.add_picture(caminho_img, width=Inches(6.2))

    # Amostra
    doc.add_heading("Amostra das Mensagens", level=2)
    df_amostra = df.head(mensagens_relatorio)

    for _, row in df_amostra.iterrows():
        texto_msg = (
            f"Data: {row.get('data', '')} {row.get('hora', '')}\n"
            f"Remetente: {row.get('remetente', '')}\n"
            f"Mensagem: {row.get('mensagem', '')}\n"
            f"Categoria: {row.get('categoria', '')}\n"
        )
        doc.add_paragraph(texto_msg)
        doc.add_paragraph("-" * 40)

    doc.save(filename)
    return filename

