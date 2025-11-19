import pandas as pd
import plotly.express as px
import io
from docx import Document
from datetime import datetime

def gerar_grafico(df):
    contagem = df["Categoria"].value_counts().reset_index()
    contagem.columns = ["Categoria", "Quantidade"]

    fig = px.bar(contagem, x="Categoria", y="Quantidade", text="Quantidade")
    fig.update_traces(textposition="outside")
    return fig


def exportar_word(df):
    doc = Document()

    doc.add_heading("Relatório de Mensagens", level=1)
    doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    tabela = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = tabela.rows[0].cells

    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col

    for _, row in df.iterrows():
        linha = tabela.add_row().cells
        for i, val in enumerate(row):
            linha[i].text = str(val)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
